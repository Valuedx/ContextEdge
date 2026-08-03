"""Phase 4d: DOCX adapter, with tracked changes handled explicitly."""

from __future__ import annotations

import io

import pytest

from contextedge.services.artifact_extraction_service import (
    _document_coverage,
    extract_artifact_text,
)
from contextedge.services.documents.base import (
    ELEMENT_HEADING,
    ELEMENT_PARAGRAPH,
    ELEMENT_TABLE,
)
from contextedge.services.documents.registry import get_parser

pytest.importorskip("docx", reason="document extras not installed")

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _doc():
    import docx

    return docx.Document()


def _bytes(document) -> bytes:
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _tracked_change_doc(deleted: str, inserted: str) -> bytes:
    """A document mid-revision, with real <w:del> / <w:ins> markup."""
    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    d = docx.Document()
    d.add_heading("5. Certificate Renewal", level=1)
    d.add_heading("5.3 Rollback", level=2)
    p = d.add_paragraph()

    def _run(text: str, wrapper: str | None):
        holder = p._p
        if wrapper:
            wrap = OxmlElement(wrapper)
            wrap.set(qn("w:id"), "1")
            wrap.set(qn("w:author"), "reviewer")
            wrap.set(qn("w:date"), "2026-01-01T00:00:00Z")
            p._p.append(wrap)
            holder = wrap
        run = OxmlElement("w:r")
        node = OxmlElement("w:delText" if wrapper == "w:del" else "w:t")
        node.text = text
        node.set(qn("xml:space"), "preserve")
        run.append(node)
        holder.append(run)

    _run(deleted, "w:del")
    _run(inserted, "w:ins")
    d.add_paragraph("Verify a test user can authenticate.")
    return _bytes(d)


# --- the reason this adapter is not a wrapper --------------------------------


def test_deleted_text_is_dropped_and_inserted_text_is_kept():
    """A naive extractor emits both, yielding "restart the database, and
    also reload without restarting" — a procedural contradiction that
    exists nowhere in the document as Word shows it. In a system whose
    job is surfacing SOP-vs-practice conflicts, that fabricated one
    reaches a reviewer as genuine.
    """
    data = _tracked_change_doc(
        "Restart the database service.",
        "Reload the certificate without restarting.",
    )
    parser = get_parser(filename="sop.docx", mime_type=None)
    parsed = parser.parse(data, filename="sop.docx")

    body = "\n".join(e.text for e in parsed.elements)
    assert "Reload the certificate without restarting." in body
    assert "Restart the database service." not in body


def test_tracked_change_volume_is_reported():
    """Dropping deletions is correct, not a degradation — so it is not a
    warning. But a document under heavy revision is a signal about
    whether it should be trusted as current."""
    data = _tracked_change_doc("old text", "new text")
    parsed = get_parser(filename="x.docx", mime_type=None).parse(data)
    tracked = parsed.metadata["tracked_changes"]
    assert tracked["deleted_runs"] >= 1
    assert tracked["inserted_runs"] >= 1


def test_a_clean_document_reports_no_tracked_changes():
    d = _doc()
    d.add_paragraph("Nothing was revised here.")
    parsed = get_parser(filename="x.docx", mime_type=None).parse(_bytes(d))
    assert parsed.metadata["tracked_changes"] == {
        "deleted_runs": 0,
        "inserted_runs": 0,
    }


# --- structure ---------------------------------------------------------------


def test_headings_nest_by_word_outline_level():
    """More reliable than the PDF adapter's font-size inference — Word
    states the level outright."""
    data = _tracked_change_doc("a", "b")
    parsed = get_parser(filename="x.docx", mime_type=None).parse(data)
    deepest = [e for e in parsed.elements if e.element_type == ELEMENT_HEADING][-1]
    assert deepest.section_path == ["5. Certificate Renewal", "5.3 Rollback"]


def test_tables_interleave_in_reading_order():
    """python-docx exposes paragraphs and tables as separate lists, which
    loses their interleaving — a table would land after every paragraph
    regardless of where it sits, detaching it from its section."""
    d = _doc()
    d.add_heading("3. Procedure", level=1)
    d.add_paragraph("Step one prose.")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Error"
    table.cell(0, 1).text = "Action"
    table.cell(1, 0).text = "401"
    table.cell(1, 1).text = "Rotate"
    d.add_heading("4. Rollback", level=1)
    d.add_paragraph("Restore the backup.")

    parsed = get_parser(filename="x.docx", mime_type=None).parse(_bytes(d))
    kinds = [e.element_type for e in parsed.elements]
    assert kinds == [
        ELEMENT_HEADING,
        ELEMENT_PARAGRAPH,
        ELEMENT_TABLE,
        ELEMENT_HEADING,
        ELEMENT_PARAGRAPH,
    ]
    table_element = next(e for e in parsed.elements if e.element_type == ELEMENT_TABLE)
    assert table_element.section_path == ["3. Procedure"]
    assert table_element.structured_content["rows"][0][0] == "Error"


def test_word_documents_have_no_page_numbers_and_say_so():
    """Word has no page concept before rendering. Citations degrade to
    "§ 5.3 Rollback" rather than inventing a page."""
    parsed = get_parser(filename="x.docx", mime_type=None).parse(
        _tracked_change_doc("a", "b")
    )
    assert parsed.page_count == 0
    assert all(e.page_number is None for e in parsed.elements)
    assert parsed.elements[-1].section_ref().startswith("§ ")


# --- integration -------------------------------------------------------------


def test_docx_reaches_the_artifact_pipeline():
    d = _doc()
    d.add_heading("Resolution:", level=1)
    d.add_paragraph("Shorten the username.")
    result = extract_artifact_text(
        filename="sop.docx", mime_type=DOCX_MIME, data=_bytes(d)
    )
    assert result.status == "completed"
    assert result.parser_type == "docx_native"
    assert "# Resolution:" in result.text
    assert "Shorten the username." in result.text


def test_page_less_formats_are_not_reported_as_zero_coverage():
    """page_count == 0 made every DOCX report coverage 0.0 — "we read
    nothing" for a perfectly parsed document, which downstream
    completeness assessment believes."""
    assert _document_coverage(0, 0, element_count=5) == 1.0
    assert _document_coverage(0, 0, element_count=0) == 0.0
    # Paginated behaviour unchanged.
    assert _document_coverage(4, 1) == 0.75


def test_docx_earns_the_document_body_budget():
    from contextedge.services.artifact_extraction_service import DOCUMENT_PARSER_TYPES

    assert "docx_native" in DOCUMENT_PARSER_TYPES


def test_docx_is_claimed_by_extension_and_mime():
    assert get_parser(filename="a.docx", mime_type=None) is not None
    assert get_parser(filename=None, mime_type=DOCX_MIME) is not None
    assert get_parser(filename="a.DOCX", mime_type=None) is not None
    # And does not steal PDFs.
    assert get_parser(filename="a.pdf", mime_type=None).name == "pdf_native"


def test_corrupt_docx_reports_a_parse_failure():
    result = extract_artifact_text(
        filename="broken.docx", mime_type=DOCX_MIME, data=b"PK\x03\x04 not really"
    )
    assert result.status == "failed"
    assert result.parser_type == "docx_native"
    assert result.error.startswith("document_parse_failed")
